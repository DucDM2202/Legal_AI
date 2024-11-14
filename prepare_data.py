import os
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.documents import Document
import pickle
from typing import List
from docx import Document as DocxDocument


def load_docx(file_path) -> List[Document]:
    """Load documents from file .docx and process paragraphs and tables."""
    doc = DocxDocument(file_path)
    content = []
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph element
            para = element.xpath(".//w:t")
            if para:
                text = ''.join([node.text for node in para if node.text])
                if text.strip():
                    content.append(text.strip())  # Each paragraph as a chunk
        elif element.tag.endswith('tbl'):  # Table element
            table_idx = [idx for idx, tbl in enumerate(doc.tables) if tbl._element == element]
            if table_idx:
                table = doc.tables[table_idx[0]]
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        content.append(' | '.join(row_data))  # Combine row cells into a single chunk
                        
    return [Document(page_content=c) for c in content]


def save_2_pickle(documents: List[Document], file_path: str):
    """Save the list of Document objects to a pickle file."""
    with open(file_path, "wb") as file:
        pickle.dump(documents, file)


def prepare(directory_path: str, file_destination_path: str):
    """Prepare and save data from all .docx files in a directory into a single pickle."""
    all_documents = []
    
    # Iterate over all files in the directory
    for filename in os.listdir(directory_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(directory_path, filename)
            documents = load_docx(file_path)
            all_documents.extend(documents)
            print(f"Loaded data from {file_path}")
    
    # Save all collected documents into a single pickle file
    save_2_pickle(all_documents, file_destination_path)
    print(f"Saved to {file_destination_path}")
def load_data(file_path: str) -> List[Document]:
    with open(file_path, "rb") as file:
        documents = pickle.load(file)
    return documents

if __name__ == "__main__":
    prepare(
        "./data/dan_su/",
        "./datapickle/dan_su.pkl"
    )
    prepare(
        "./data/hien_phap/",
        "./datapickle/hien_phap.pkl"
    )
    prepare(
        "./data/hinh_su/",
        "./datapickle/hinh_su.pkl"
    )
    prepare(
        "./data/dat_dai/",
        "./datapickle/dat_dai.pkl"
    )
    prepare(
        "./data/hien_phap/",
        "./datapickle/hien_phap.pkl"
    )
    prepare(
        "./data/hon_nhan_gia_dinh/",
        "./datapickle/hon_nhan_gia_dinh.pkl"
    )
    prepare(
        "./data/lao_dong/",
        "./datapickle/lao_dong.pkl"
    )
    prepare(
        "./data/tai_chinh/",
        "./datapickle/tai_chinh.pkl"
    )


    